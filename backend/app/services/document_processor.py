import os
import uuid
import magic
import re
import aiofiles
import asyncio
import logging
import time
from datetime import datetime, timezone
from typing import List, Tuple, Optional, Dict, Any, Union
from pathlib import Path
from fastapi import UploadFile
from app.models.document import Document, UploadedFileInfo, UploadValidationError
from app.core.config import settings
from app.services.exceptions import (
    DocumentProcessingError, UnsupportedFileTypeError, FileAccessError,
    PDFExtractionError, DOCXExtractionError, TextExtractionError,
    CorruptedFileError, PasswordProtectedError, FileSizeError,
    MetadataExtractionError, EmptyDocumentError
)

# Document processing libraries
import PyPDF2
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.table import Table
import markdown
from bs4 import BeautifulSoup
import chardet
import pandas as pd

# Text processing
import nltk
from langdetect import detect, DetectorFactory

# Set seed for consistent language detection
DetectorFactory.seed = 0

# Configure logging
logger = logging.getLogger(__name__)


class DocumentMetadata:
    """Class to hold document metadata information."""
    
    def __init__(self, file_path: str, file_size: int = 0):
        self.file_path = file_path
        self.file_size = file_size
        self.word_count = 0
        self.character_count = 0
        self.line_count = 0
        self.page_count = 0
        self.language = "unknown"
        self.encoding = "unknown"
        self.creation_date: Optional[datetime] = None
        self.modification_date: Optional[datetime] = None
        self.file_type = "unknown"
        self.mime_type = "unknown"
        self.has_tables = False
        self.table_count = 0
        self.has_images = False
        self.image_count = 0
        self.is_password_protected = False
        self.extraction_method = "unknown"
        self.processing_time = 0.0
        self.errors: List[str] = []
        self.warnings: List[str] = []
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary format."""
        return {
            "file_path": self.file_path,
            "file_size": self.file_size,
            "word_count": self.word_count,
            "character_count": self.character_count,
            "line_count": self.line_count,
            "page_count": self.page_count,
            "language": self.language,
            "encoding": self.encoding,
            "creation_date": self.creation_date.isoformat() if self.creation_date else None,
            "modification_date": self.modification_date.isoformat() if self.modification_date else None,
            "file_type": self.file_type,
            "mime_type": self.mime_type,
            "has_tables": self.has_tables,
            "table_count": self.table_count,
            "has_images": self.has_images,
            "image_count": self.image_count,
            "is_password_protected": self.is_password_protected,
            "extraction_method": self.extraction_method,
            "processing_time": self.processing_time,
            "errors": self.errors,
            "warnings": self.warnings
        }


class DocumentProcessor:
    """Enhanced document processor with comprehensive text extraction and metadata analysis."""
    
    def __init__(self):
        self.documents_folder = settings.documents_folder
        self._vector_store = None
        self._llm_service = None
        self.max_file_size = getattr(settings, 'max_file_size_bytes', 50 * 1024 * 1024)  # 50MB default
        self.supported_encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'ascii']
        
        # Initialize NLTK data if needed
        self._ensure_nltk_data()
    
    async def _get_vector_store(self):
        """Get or initialize vector store instance."""
        if self._vector_store is None:
            from app.services.vector_store import get_vector_store
            self._vector_store = await get_vector_store()
        return self._vector_store
    
    def _get_llm_service(self):
        """Get or initialize LLM service instance."""
        if self._llm_service is None:
            from app.services.llm_service import LLMService
            self._llm_service = LLMService()
        return self._llm_service
    
    def _ensure_nltk_data(self):
        """Ensure required NLTK data is downloaded."""
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            try:
                nltk.download('punkt', quiet=True)
            except Exception as e:
                logger.warning(f"Failed to download NLTK punkt tokenizer: {e}")
        
        try:
            nltk.data.find('tokenizers/punkt_tab')
        except LookupError:
            try:
                nltk.download('punkt_tab', quiet=True)
            except Exception as e:
                logger.warning(f"Failed to download NLTK punkt_tab tokenizer: {e}")
    
    async def scan_documents_folder(self) -> List[Document]:
        """Scan the documents folder and return list of documents"""
        documents = []
        
        if not os.path.exists(self.documents_folder):
            return documents
        
        for filename in os.listdir(self.documents_folder):
            if self._is_supported_file(filename):
                file_path = os.path.join(self.documents_folder, filename)
                
                # Extract original filename if this was an uploaded file
                display_filename = self._extract_original_filename(filename)
                
                document = Document(
                    id=str(uuid.uuid4()),
                    filename=display_filename,
                    file_path=file_path,
                    status="ready",
                    created_at=datetime.fromtimestamp(os.path.getctime(file_path)),
                    chunk_count=0
                )
                documents.append(document)
        
        return documents
    
    def _extract_original_filename(self, filename: str) -> str:
        """Extract original filename from uploaded file naming convention"""
        # Check if filename follows our upload naming pattern: YYYYMMDD_HHMMSS_XXXXXXXX_originalname.ext
        import re
        pattern = r'^\d{8}_\d{6}_[a-f0-9]{8}_(.+)$'
        match = re.match(pattern, filename)
        
        if match:
            return match.group(1)  # Return the original filename part
        else:
            return filename  # Return as-is for non-uploaded files
    
    def _is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        supported_extensions = ['.pdf', '.docx', '.txt', '.md']
        return any(filename.lower().endswith(ext) for ext in supported_extensions)
    
    async def extract_text_from_file(self, file_path: str, extract_metadata: bool = True) -> Union[str, Tuple[str, DocumentMetadata]]:
        """Extract text from a file based on its type with enhanced error handling.
        
        Args:
            file_path: Path to the file to extract text from
            extract_metadata: Whether to also extract and return metadata
        
        Returns:
            If extract_metadata is False: extracted text string
            If extract_metadata is True: tuple of (text, metadata)
        
        Raises:
            UnsupportedFileTypeError: If file type is not supported
            FileAccessError: If file cannot be accessed
            Various specific extraction errors based on file type
        """
        start_time = time.time()
        
        # Validate file exists and is accessible
        if not os.path.exists(file_path):
            raise FileAccessError(file_path, FileNotFoundError(f"File not found: {file_path}"))
        
        if not os.path.isfile(file_path):
            raise FileAccessError(file_path, ValueError(f"Path is not a file: {file_path}"))
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            raise FileSizeError(file_path, file_size, self.max_file_size)
        
        # Initialize metadata if requested
        metadata = None
        if extract_metadata:
            metadata = DocumentMetadata(file_path, file_size)
            metadata.creation_date = datetime.fromtimestamp(os.path.getctime(file_path))
            metadata.modification_date = datetime.fromtimestamp(os.path.getmtime(file_path))
        
        filename = os.path.basename(file_path).lower()
        text = ""
        
        try:
            if filename.endswith('.pdf'):
                text = await self._extract_from_pdf_enhanced(file_path, metadata)
            elif filename.endswith('.docx'):
                text = await self._extract_from_docx_enhanced(file_path, metadata)
            elif filename.endswith('.txt'):
                text = await self._extract_from_txt_enhanced(file_path, metadata)
            elif filename.endswith('.md'):
                text = await self._extract_from_markdown_enhanced(file_path, metadata)
            else:
                supported_types = ['.pdf', '.docx', '.txt', '.md']
                raise UnsupportedFileTypeError(file_path, os.path.splitext(filename)[1], supported_types)
            
            # Post-process text
            text = self._clean_and_normalize_text(text)
            
            # Update metadata if requested
            if extract_metadata:
                metadata.processing_time = time.time() - start_time
                self._analyze_text_metadata(text, metadata)
                
                # Check if document is empty
                if not text.strip():
                    metadata.warnings.append("Document contains no extractable text")
                    
                return text, metadata
            
            return text
            
        except DocumentProcessingError:
            # Re-raise our custom exceptions
            raise
        except Exception as e:
            logger.error(f"Unexpected error extracting text from {file_path}: {e}")
            if extract_metadata:
                metadata.processing_time = time.time() - start_time
                metadata.errors.append(f"Unexpected error: {str(e)}")
                return "", metadata
            raise DocumentProcessingError(f"Unexpected error during text extraction: {str(e)}", file_path)
    
    async def _extract_from_pdf_enhanced(self, file_path: str, metadata: Optional[DocumentMetadata] = None) -> str:
        """Enhanced PDF text extraction with multiple fallback methods."""
        text = ""
        extraction_method = "unknown"
        
        try:
            # Method 1: Try PyMuPDF first (more robust)
            try:
                doc = fitz.open(file_path)
                if doc.is_encrypted:
                    doc.close()
                    raise PasswordProtectedError(file_path)
                
                page_texts = []
                for page_num in range(len(doc)):
                    try:
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                        page_texts.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1} from PDF {file_path}: {e}")
                        if metadata:
                            metadata.warnings.append(f"Failed to extract page {page_num + 1}: {str(e)}")
                        page_texts.append("")  # Add empty page to maintain structure
                
                text = "\n\n".join(page_texts)
                extraction_method = "PyMuPDF"
                
                if metadata:
                    metadata.page_count = len(doc)
                    metadata.extraction_method = extraction_method
                    # Check for images
                    image_count = 0
                    for page_num in range(len(doc)):
                        page = doc.load_page(page_num)
                        image_list = page.get_images()
                        image_count += len(image_list)
                    metadata.has_images = image_count > 0
                    metadata.image_count = image_count
                
                doc.close()
                
            except PasswordProtectedError:
                raise
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed for {file_path}: {e}")
                
                # Method 2: Fallback to PyPDF2
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        
                        if pdf_reader.is_encrypted:
                            raise PasswordProtectedError(file_path)
                        
                        page_texts = []
                        for page_num, page in enumerate(pdf_reader.pages):
                            try:
                                page_text = page.extract_text()
                                page_texts.append(page_text)
                            except Exception as page_error:
                                logger.warning(f"Failed to extract page {page_num + 1} with PyPDF2: {page_error}")
                                if metadata:
                                    metadata.warnings.append(f"Failed to extract page {page_num + 1} with PyPDF2: {str(page_error)}")
                                page_texts.append("")  # Add empty page to maintain structure
                        
                        text = "\n\n".join(page_texts)
                        extraction_method = "PyPDF2"
                        
                        if metadata:
                            metadata.page_count = len(pdf_reader.pages)
                            metadata.extraction_method = extraction_method
                            
                except Exception as e2:
                    logger.error(f"Both PDF extraction methods failed for {file_path}: PyMuPDF: {e}, PyPDF2: {e2}")
                    raise PDFExtractionError(file_path, original_error=e2)
        
        except PasswordProtectedError:
            if metadata:
                metadata.is_password_protected = True
                metadata.errors.append("File is password-protected")
            raise
        except PDFExtractionError:
            raise
        except Exception as e:
            raise PDFExtractionError(file_path, original_error=e)
        
        # Validate extracted text
        if not text.strip():
            if metadata:
                metadata.warnings.append("No text could be extracted from PDF")
            logger.warning(f"No text extracted from PDF: {file_path}")
        
        return text
    
    async def _extract_from_docx_enhanced(self, file_path: str, metadata: Optional[DocumentMetadata] = None) -> str:
        """Enhanced DOCX text extraction with table support and formatting preservation."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            table_count = 0
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_count += 1
                table_text = self._extract_table_text(table)
                if table_text.strip():
                    text_parts.append(f"\n--- Table {table_count} ---\n{table_text}\n--- End Table ---\n")
            
            text = "\n".join(text_parts)
            
            if metadata:
                metadata.file_type = "DOCX"
                metadata.extraction_method = "python-docx"
                metadata.has_tables = table_count > 0
                metadata.table_count = table_count
                
                # Try to get additional document properties
                try:
                    core_props = doc.core_properties
                    if hasattr(core_props, 'created') and core_props.created:
                        metadata.creation_date = core_props.created
                    if hasattr(core_props, 'modified') and core_props.modified:
                        metadata.modification_date = core_props.modified
                except Exception as e:
                    metadata.warnings.append(f"Could not extract document properties: {str(e)}")
            
            return text
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise DOCXExtractionError(file_path, original_error=e)
    
    def _extract_table_text(self, table: Table) -> str:
        """Extract text from a DOCX table with proper formatting."""
        try:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Clean cell text and handle line breaks
                    cell_text = cell.text.strip().replace('\n', ' ').replace('\r', '')
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            # Convert to DataFrame for better formatting
            if table_data:
                df = pd.DataFrame(table_data[1:], columns=table_data[0] if len(table_data) > 1 else None)
                return df.to_string(index=False, na_rep='')
            return ""
            
        except Exception as e:
            logger.warning(f"Failed to extract table text: {e}")
            # Fallback to simple text extraction
            text_parts = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
            return "\n".join(text_parts)
    
    async def _extract_from_txt_enhanced(self, file_path: str, metadata: Optional[DocumentMetadata] = None) -> str:
        """Enhanced TXT extraction with encoding detection and error recovery."""
        text = ""
        detected_encoding = "unknown"
        
        try:
            # First, try to detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                
            # Use chardet for encoding detection
            encoding_result = chardet.detect(raw_data)
            detected_encoding = encoding_result.get('encoding', 'utf-8')
            confidence = encoding_result.get('confidence', 0.0)
            
            if metadata:
                metadata.encoding = detected_encoding
                if confidence < 0.7:
                    metadata.warnings.append(f"Low confidence ({confidence:.2f}) in encoding detection")
            
            # Try detected encoding first
            encodings_to_try = [detected_encoding] + [enc for enc in self.supported_encodings if enc != detected_encoding]
            
            for encoding in encodings_to_try:
                if not encoding:
                    continue
                    
                try:
                    text = raw_data.decode(encoding)
                    if metadata:
                        metadata.encoding = encoding
                        metadata.extraction_method = f"Direct decode ({encoding})"
                    break
                except (UnicodeDecodeError, LookupError) as e:
                    logger.debug(f"Failed to decode {file_path} with {encoding}: {e}")
                    continue
            
            if not text:
                # Last resort: decode with errors='ignore'
                try:
                    text = raw_data.decode('utf-8', errors='ignore')
                    if metadata:
                        metadata.encoding = "utf-8 (with errors ignored)"
                        metadata.extraction_method = "UTF-8 with error recovery"
                        metadata.warnings.append("Some characters may have been lost during encoding recovery")
                except Exception as e:
                    raise TextExtractionError(file_path, encoding="utf-8 (error recovery)", original_error=e)
            
            if metadata:
                metadata.file_type = "TXT"
                metadata.line_count = len(text.splitlines())
            
            return text
            
        except TextExtractionError:
            raise
        except Exception as e:
            logger.error(f"Failed to extract text from TXT file {file_path}: {e}")
            raise TextExtractionError(file_path, encoding=detected_encoding, original_error=e)
    
    async def _extract_from_markdown_enhanced(self, file_path: str, metadata: Optional[DocumentMetadata] = None) -> str:
        """Enhanced Markdown extraction with proper text conversion and metadata."""
        try:
            # First extract raw markdown content with encoding detection
            raw_text = await self._extract_from_txt_enhanced(file_path, metadata)
            
            # Convert markdown to HTML, then to plain text
            html_content = markdown.markdown(raw_text, extensions=['tables', 'fenced_code', 'toc'])
            
            # Use BeautifulSoup to extract clean text from HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text while preserving some structure
            text_parts = []
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'td', 'th']):
                text = element.get_text().strip()
                if text:
                    # Add extra spacing for headers
                    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        text_parts.append(f"\n{text}\n")
                    else:
                        text_parts.append(text)
            
            # Handle code blocks separately
            for code_block in soup.find_all(['pre', 'code']):
                code_text = code_block.get_text().strip()
                if code_text and len(code_text) > 20:  # Only include substantial code blocks
                    text_parts.append(f"\n--- Code Block ---\n{code_text}\n--- End Code ---\n")
            
            final_text = "\n".join(text_parts)
            
            if metadata:
                metadata.file_type = "Markdown"
                metadata.extraction_method = "markdown + BeautifulSoup"
                
                # Count tables
                table_count = len(soup.find_all('table'))
                metadata.has_tables = table_count > 0
                metadata.table_count = table_count
                
                # Check for code blocks
                code_blocks = len(soup.find_all(['pre', 'code']))
                if code_blocks > 0:
                    metadata.warnings.append(f"Document contains {code_blocks} code blocks")
            
            return final_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from Markdown file {file_path}: {e}")
            raise TextExtractionError(file_path, original_error=e)
    
    async def process_all_documents(self) -> dict:
        """Process all documents in the folder with enhanced reporting."""
        documents = await self.scan_documents_folder()
        processed_count = 0
        failed_documents = []
        processing_stats = {
            "total_time": 0.0,
            "total_characters": 0,
            "total_words": 0,
            "by_file_type": {}
        }
        
        for document in documents:
            start_time = time.time()
            try:
                text, metadata = await self.extract_text_from_file(document.file_path, extract_metadata=True)
                
                if text.strip():
                    processed_count += 1
                    
                    # Update statistics
                    processing_time = time.time() - start_time
                    processing_stats["total_time"] += processing_time
                    processing_stats["total_characters"] += metadata.character_count
                    processing_stats["total_words"] += metadata.word_count
                    
                    file_type = metadata.file_type or "unknown"
                    if file_type not in processing_stats["by_file_type"]:
                        processing_stats["by_file_type"][file_type] = {"count": 0, "words": 0, "chars": 0}
                    
                    processing_stats["by_file_type"][file_type]["count"] += 1
                    processing_stats["by_file_type"][file_type]["words"] += metadata.word_count
                    processing_stats["by_file_type"][file_type]["chars"] += metadata.character_count
                    
                    # Add document to vector store
                    try:
                        vector_store = await self._get_vector_store()
                        document_id = str(uuid.uuid4())
                        
                        # Add document metadata
                        doc_metadata = {
                            "file_path": document.file_path,
                            "file_size": metadata.file_size,
                            "word_count": metadata.word_count,
                            "page_count": metadata.page_count,
                            "language": metadata.language,
                            "file_type": metadata.file_type,
                            "processed_at": datetime.now(timezone.utc).isoformat()
                        }
                        
                        success = await vector_store.add_document(
                            document_id=document_id,
                            document_name=document.filename,
                            text=text,
                            metadata=doc_metadata
                        )
                        
                        if success:
                            logger.info(f"Successfully added {document.filename} to vector store")
                        else:
                            logger.warning(f"Failed to add {document.filename} to vector store")
                            
                    except Exception as e:
                        logger.error(f"Error adding {document.filename} to vector store: {e}")
                        # Continue processing other documents even if vector store fails
                else:
                    failed_documents.append({
                        "filename": document.filename,
                        "error": "No text extracted",
                        "warnings": metadata.warnings if metadata else []
                    })
                    
            except Exception as e:
                logger.error(f"Error processing {document.filename}: {e}")
                failed_documents.append({
                    "filename": document.filename,
                    "error": str(e),
                    "error_type": type(e).__name__
                })
        
        return {
            "total_documents": len(documents),
            "processed_count": processed_count,
            "failed_count": len(failed_documents),
            "failed_documents": failed_documents,
            "processing_stats": processing_stats
        }
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent directory traversal and other security issues"""
        # Remove directory components
        filename = os.path.basename(filename)
        
        # Remove or replace dangerous characters
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        
        # Remove leading/trailing dots and spaces
        filename = filename.strip('. ')
        
        # Ensure filename is not empty and not reserved
        if not filename or filename.lower() in ['con', 'prn', 'aux', 'nul'] + [f'com{i}' for i in range(1, 10)] + [f'lpt{i}' for i in range(1, 10)]:
            filename = f"file_{uuid.uuid4().hex[:8]}"
        
        return filename
    
    def _generate_unique_filename(self, original_filename: str) -> str:
        """Generate a unique filename with timestamp and UUID to prevent conflicts"""
        sanitized = self._sanitize_filename(original_filename)
        name, ext = os.path.splitext(sanitized)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = uuid.uuid4().hex[:8]
        return f"{timestamp}_{unique_id}_{name}{ext}"
    
    def _validate_file_type(self, file_content: bytes, filename: str) -> Tuple[bool, str, str]:
        """Validate file type using both extension and MIME type detection"""
        # Check extension
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext not in settings.allowed_extensions:
            return False, "Invalid file extension", ""
        
        # Detect MIME type from content
        try:
            mime_type = magic.from_buffer(file_content, mime=True)
        except Exception as e:
            return False, f"Failed to detect file type: {str(e)}", ""
        
        # Check if MIME type is allowed
        if mime_type not in settings.allowed_file_types:
            # Handle some common MIME type variations
            mime_mappings = {
                "text/x-markdown": "text/markdown",
                "application/x-empty": "text/plain"  # Empty files
            }
            mime_type = mime_mappings.get(mime_type, mime_type)
            
            if mime_type not in settings.allowed_file_types:
                return False, f"File type not allowed: {mime_type}", mime_type
        
        return True, "Valid", mime_type
    
    def _validate_file_size(self, file_size: int) -> Tuple[bool, str]:
        """Validate file size against configured limits"""
        if file_size > settings.max_file_size_bytes:
            return False, f"File size ({file_size / 1024 / 1024:.2f}MB) exceeds maximum allowed size ({settings.max_file_size_mb}MB)"
        
        if file_size == 0:
            return False, "File is empty"
        
        return True, "Valid"
    
    async def validate_upload_file(self, upload_file: UploadFile) -> Tuple[bool, List[UploadValidationError]]:
        """Validate an uploaded file for size, type, and content"""
        errors = []
        
        if not upload_file.filename:
            errors.append(UploadValidationError(
                filename="unknown",
                error_type="filename",
                message="No filename provided"
            ))
            return False, errors
        
        # Read file content for validation
        try:
            file_content = await upload_file.read()
            await upload_file.seek(0)  # Reset file pointer
        except Exception as e:
            errors.append(UploadValidationError(
                filename=upload_file.filename,
                error_type="read_error",
                message=f"Failed to read file: {str(e)}"
            ))
            return False, errors
        
        # Validate file size
        file_size = len(file_content)
        size_valid, size_message = self._validate_file_size(file_size)
        if not size_valid:
            errors.append(UploadValidationError(
                filename=upload_file.filename,
                error_type="size",
                message=size_message
            ))
        
        # Validate file type
        type_valid, type_message, mime_type = self._validate_file_type(file_content, upload_file.filename)
        if not type_valid:
            errors.append(UploadValidationError(
                filename=upload_file.filename,
                error_type="type",
                message=type_message
            ))
        
        return len(errors) == 0, errors
    
    async def save_uploaded_file(self, upload_file: UploadFile) -> UploadedFileInfo:
        """Save an uploaded file to the documents folder"""
        # Generate unique filename
        unique_filename = self._generate_unique_filename(upload_file.filename)
        file_path = os.path.join(self.documents_folder, unique_filename)
        
        # Ensure documents folder exists
        os.makedirs(self.documents_folder, exist_ok=True)
        
        # Read file content
        file_content = await upload_file.read()
        
        # Get MIME type
        try:
            mime_type = magic.from_buffer(file_content, mime=True)
        except Exception:
            # Fallback to content type from upload
            mime_type = upload_file.content_type or "application/octet-stream"
        
        # Save file asynchronously
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        # Create file info
        file_info = UploadedFileInfo(
            original_filename=upload_file.filename,
            saved_filename=unique_filename,
            file_size=len(file_content),
            content_type=mime_type,
            upload_timestamp=datetime.now(),
            document_id=str(uuid.uuid4())
        )
        
        return file_info
    
    async def process_file_uploads(self, upload_files: List[UploadFile]) -> Tuple[List[UploadedFileInfo], List[UploadValidationError]]:
        """Process multiple file uploads with validation"""
        uploaded_files = []
        all_errors = []
        
        # Validate number of files
        if len(upload_files) > settings.max_files_per_upload:
            all_errors.append(UploadValidationError(
                filename="multiple",
                error_type="count",
                message=f"Too many files. Maximum {settings.max_files_per_upload} files allowed per upload."
            ))
            return uploaded_files, all_errors
        
        # Process each file
        for upload_file in upload_files:
            try:
                # Validate file
                is_valid, validation_errors = await self.validate_upload_file(upload_file)
                
                if is_valid:
                    # Save file
                    file_info = await self.save_uploaded_file(upload_file)
                    uploaded_files.append(file_info)
                else:
                    all_errors.extend(validation_errors)
                    
            except Exception as e:
                all_errors.append(UploadValidationError(
                    filename=upload_file.filename or "unknown",
                    error_type="processing",
                    message=f"Unexpected error processing file: {str(e)}"
                ))
        
        return uploaded_files, all_errors
    
    async def get_upload_statistics(self) -> dict:
        """Get statistics about uploaded files"""
        if not os.path.exists(self.documents_folder):
            return {"total_files": 0, "uploaded_files": 0, "original_files": 0}
        
        total_files = 0
        uploaded_files = 0
        original_files = 0
        
        for filename in os.listdir(self.documents_folder):
            if self._is_supported_file(filename):
                total_files += 1
                # Check if it's an uploaded file based on naming pattern
                if re.match(r'^\d{8}_\d{6}_[a-f0-9]{8}_', filename):
                    uploaded_files += 1
                else:
                    original_files += 1
        
        return {
            "total_files": total_files,
            "uploaded_files": uploaded_files,
            "original_files": original_files
        }
    
    def _clean_and_normalize_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return text
        
        # Remove excessive whitespace while preserving structure
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Clean each line
            cleaned_line = re.sub(r'\s+', ' ', line.strip())
            cleaned_lines.append(cleaned_line)
        
        # Remove excessive empty lines (keep max 2 consecutive)
        result_lines = []
        empty_count = 0
        
        for line in cleaned_lines:
            if not line:
                empty_count += 1
                if empty_count <= 2:
                    result_lines.append(line)
            else:
                empty_count = 0
                result_lines.append(line)
        
        # Join and clean up final text
        final_text = '\n'.join(result_lines)
        
        # Remove common OCR artifacts and weird characters
        final_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x9F]', '', final_text)
        
        # Normalize unicode
        try:
            import unicodedata
            final_text = unicodedata.normalize('NFKC', final_text)
        except Exception:
            pass  # Skip normalization if it fails
        
        return final_text.strip()
    
    def _analyze_text_metadata(self, text: str, metadata: DocumentMetadata):
        """Analyze text content and update metadata."""
        if not text:
            return
        
        try:
            # Basic counts
            metadata.character_count = len(text)
            metadata.line_count = len(text.splitlines())
            
            # Word count using NLTK if available
            try:
                from nltk.tokenize import word_tokenize
                words = word_tokenize(text)
                metadata.word_count = len([word for word in words if word.isalnum()])
            except Exception:
                # Fallback to simple split
                metadata.word_count = len(text.split())
            
            # Language detection
            try:
                if len(text.strip()) > 50:  # Need sufficient text for detection
                    detected_lang = detect(text[:1000])  # Use first 1000 chars for speed
                    metadata.language = detected_lang
            except Exception as e:
                metadata.warnings.append(f"Language detection failed: {str(e)}")
                metadata.language = "unknown"
        
        except Exception as e:
            logger.warning(f"Failed to analyze text metadata: {e}")
            metadata.warnings.append(f"Text analysis failed: {str(e)}")
    
    async def get_document_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract comprehensive metadata from a document without full text extraction."""
        try:
            _, metadata = await self.extract_text_from_file(file_path, extract_metadata=True)
            return metadata
        except Exception as e:
            # Return basic metadata even if extraction fails
            metadata = DocumentMetadata(file_path)
            metadata.errors.append(f"Metadata extraction failed: {str(e)}")
            
            # Still try to get basic file info
            try:
                stat = os.stat(file_path)
                metadata.file_size = stat.st_size
                metadata.creation_date = datetime.fromtimestamp(stat.st_ctime)
                metadata.modification_date = datetime.fromtimestamp(stat.st_mtime)
                metadata.file_type = os.path.splitext(file_path)[1][1:].upper()
            except Exception:
                pass
            
            return metadata
    
    async def validate_document_integrity(self, file_path: str) -> Tuple[bool, List[str]]:
        """Validate document integrity and readability."""
        issues = []
        
        try:
            # Basic file checks
            if not os.path.exists(file_path):
                issues.append("File does not exist")
                return False, issues
            
            if not os.path.isfile(file_path):
                issues.append("Path is not a file")
                return False, issues
            
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                issues.append("File is empty")
                return False, issues
            
            if file_size > self.max_file_size:
                issues.append(f"File size ({file_size:,} bytes) exceeds maximum ({self.max_file_size:,} bytes)")
                return False, issues
            
            # Try to extract a small amount of text to verify readability
            try:
                text = await self.extract_text_from_file(file_path, extract_metadata=False)
                if not text.strip():
                    issues.append("No text could be extracted from document")
                elif len(text.strip()) < 10:
                    issues.append("Very little text extracted - document may be corrupted or mostly non-textual")
            except PasswordProtectedError:
                issues.append("Document is password-protected")
                return False, issues
            except Exception as e:
                issues.append(f"Text extraction failed: {str(e)}")
                return False, issues
            
            return len(issues) == 0, issues
            
        except Exception as e:
            issues.append(f"Validation failed: {str(e)}")
            return False, issues
    
    async def get_processing_statistics(self) -> Dict[str, Any]:
        """Get comprehensive processing statistics for all documents."""
        documents = await self.scan_documents_folder()
        
        if not documents:
            return {
                "total_documents": 0,
                "processing_summary": {},
                "file_type_breakdown": {},
                "size_distribution": {},
                "health_check": {"healthy": 0, "warnings": 0, "errors": 0}
            }
        
        stats = {
            "total_documents": len(documents),
            "processing_summary": {
                "processable": 0,
                "corrupted": 0,
                "password_protected": 0,
                "empty": 0,
                "oversized": 0
            },
            "file_type_breakdown": {},
            "size_distribution": {
                "small": 0,    # < 1MB
                "medium": 0,   # 1-10MB
                "large": 0,    # 10-50MB
                "xlarge": 0    # > 50MB
            },
            "health_check": {"healthy": 0, "warnings": 0, "errors": 0}
        }
        
        # Process each document for statistics
        for document in documents:
            try:
                # Get file info
                file_size = os.path.getsize(document.file_path) if os.path.exists(document.file_path) else 0
                file_ext = os.path.splitext(document.filename)[1].lower()
                
                # Update file type breakdown
                if file_ext not in stats["file_type_breakdown"]:
                    stats["file_type_breakdown"][file_ext] = 0
                stats["file_type_breakdown"][file_ext] += 1
                
                # Update size distribution
                if file_size < 1024 * 1024:  # < 1MB
                    stats["size_distribution"]["small"] += 1
                elif file_size < 10 * 1024 * 1024:  # < 10MB
                    stats["size_distribution"]["medium"] += 1
                elif file_size < 50 * 1024 * 1024:  # < 50MB
                    stats["size_distribution"]["large"] += 1
                else:
                    stats["size_distribution"]["xlarge"] += 1
                
                # Quick integrity check
                try:
                    is_healthy, issues = await self.validate_document_integrity(document.file_path)
                    
                    if is_healthy:
                        stats["processing_summary"]["processable"] += 1
                        stats["health_check"]["healthy"] += 1
                    else:
                        # Categorize the issue
                        has_password_error = any("password" in issue.lower() for issue in issues)
                        has_empty_error = any("empty" in issue.lower() for issue in issues)
                        has_size_error = any("size" in issue.lower() for issue in issues)
                        
                        if has_password_error:
                            stats["processing_summary"]["password_protected"] += 1
                        elif has_empty_error:
                            stats["processing_summary"]["empty"] += 1
                        elif has_size_error:
                            stats["processing_summary"]["oversized"] += 1
                        else:
                            stats["processing_summary"]["corrupted"] += 1
                        
                        stats["health_check"]["errors"] += 1
                        
                except Exception:
                    stats["processing_summary"]["corrupted"] += 1
                    stats["health_check"]["errors"] += 1
                    
            except Exception as e:
                logger.warning(f"Failed to analyze document {document.filename}: {e}")
                stats["health_check"]["errors"] += 1
        
        return stats